#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enhanced Resume Analysis System - Systematic Data Extraction
============================================================

Improvements for cleaner, more systematic data processing:
1. Better text preprocessing and normalization
2. Enhanced PDF extraction with fallback methods
3. Improved keyword matching with fuzzy logic
4. Systematic data validation and quality checks
5. Better handling of edge cases and errors
6. Integration with pyexcel-ods3 for ODS file support
"""

import argparse, os, re, sys, json, logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
from collections import Counter, defaultdict
import string
import unicodedata

import numpy as np
import pandas as pd

# Text extraction - multiple methods for reliability
import pdfplumber
from docx import Document
import pyexcel_ods3  # As specified in user rules

# Features & similarity
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import StandardScaler

# Additional libraries for better text processing
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
import spacy

# Setup logging for better debugging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class TextProcessor:
    """Enhanced text processing with multiple normalization strategies"""
    
    def __init__(self):
        self.stemmer = PorterStemmer()
        self.lemmatizer = WordNetLemmatizer()
        # Download required NLTK data if not present
        self._ensure_nltk_data()
        
        # Try to load spaCy model for advanced NLP
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except IOError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def _ensure_nltk_data(self):
        """Ensure required NLTK data is downloaded"""
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('corpora/stopwords')
            nltk.data.find('corpora/wordnet')
        except LookupError:
            nltk.download('punkt')
            nltk.download('stopwords')
            nltk.download('wordnet')
    
    def clean_text(self, text: str) -> str:
        """Comprehensive text cleaning and normalization"""
        if not text:
            return ""
        
        # 1. Unicode normalization
        text = unicodedata.normalize('NFKD', text)
        
        # 2. Remove PDF artifacts and special characters
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'[\r\n\t]+', ' ', text)  # Remove line breaks/tabs
        
        # 3. Remove common PDF extraction artifacts
        pdf_artifacts = [
            r'Page \d+ of \d+', r'©.*?\d{4}', r'www\..*?\.(com|org|edu)',
            r'\b[A-Z]{2,}\b(?=\s+[A-Z]{2,})',  # All caps sequences (headers)
            r'^\s*\|\s*', r'\s*\|\s*$',  # Table separators
            r'•\s*', r'◦\s*', r'▪\s*'  # Bullet points
        ]
        
        for pattern in pdf_artifacts:
            text = re.sub(pattern, ' ', text, flags=re.IGNORECASE | re.MULTILINE)
        
        # 4. Clean punctuation while preserving meaningful ones
        text = re.sub(r'[^\w\s\-\+\#\.\,\(\)\/]', ' ', text)
        
        # 5. Fix common OCR errors
        ocr_fixes = {
            r'\b0\b': 'O',  # Zero to O
            r'\bl\b': 'I',  # l to I (in contexts like 'l have')
            r'\brn\b': 'm',  # rn to m
        }
        for pattern, replacement in ocr_fixes.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # 6. Normalize spacing and trim
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def extract_structured_info(self, text: str) -> Dict[str, List[str]]:
        """Extract structured information like emails, phones, URLs"""
        info = defaultdict(list)
        
        # Email addresses
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        info['emails'] = emails
        
        # Phone numbers (multiple formats)
        phones = re.findall(r'(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})', text)
        info['phones'] = phones
        
        # URLs
        urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', text)
        info['urls'] = urls
        
        # LinkedIn profiles
        linkedin = re.findall(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
        info['linkedin'] = linkedin
        
        # GitHub profiles
        github = re.findall(r'github\.com/[\w\-]+', text, re.IGNORECASE)
        info['github'] = github
        
        return dict(info)
    
    def detect_degree_info(self, text: str) -> Dict[str, str]:
        """Detect degree level (bachelor/master/phd) and likely field (CS/CE) with strict patterns.
        Avoid false positives like 'MS Word' by requiring academic context.
        """
        t = text.lower()
        degree_level = 'unknown'
        degree_field = ''
        
        # Order matters: detect PhD first, then masters, then bachelors.
        phd_patterns = [
            r"\bph\.?d\.?\b",
            r"\bp\s*\.?\s*h\s*\.?\s*d\b",
            r"\bdoctor\s+of\s+philosophy\b",
            r"\bdphil\b",
            r"\bscd\b",
            r"\bsd\.?sc\b",
        ]
        
        # Masters: require explicit 'master of ...' or an abbreviation followed by context like 'in', '(', ',', or end.
        masters_patterns = [
            r"\bmaster(?:'s)?\s*of\s*(science|engineering|computer|electrical|software|information|technology)",
            r"\bmaster\s*s\b",
            r"\bmasterof(science|engineering)",
            r"\bm\s*\.?\s*s\s*\.?\b",
            r"\bm\.?sc\b",
            r"\bm\.?eng\b",
            r"\bm\.?e\b",
        ]
        # Explicitly ignore 'ms' when referring to Microsoft terms
        masters_negative = r"\bms\s+(word|excel|office|powerpoint|outlook|teams)\b"
        
        # Bachelors: similar strictness
        bachelors_patterns = [
            r"\bbachelor(?:'s)?\s*of\s*(science|engineering|arts|technology)",
            r"\bbachelor\s*s\b",
            r"\bbachelorof(science|engineering|arts|technology)",
            r"\bb\s*\.?\s*s\s*\.?\b",
            r"\bb\.?sc\b",
            r"\bb\.?eng\b",
            r"\bbe\b",
            r"\bb\.\s*tech\b",
            r"\bbtech\b",
            r"\bundergraduate\b",
        ]
        
        if any(re.search(p, t) for p in phd_patterns):
            degree_level = 'phd'
        else:
            # Masters check: ensure not matching Microsoft abbreviations
            if (any(re.search(p, t) for p in masters_patterns) and not re.search(masters_negative, t)):
                degree_level = 'masters'
            elif any(re.search(p, t) for p in bachelors_patterns):
                degree_level = 'bachelors'
        
        # Field detection
        cs_patterns = [r"computer\s+science", r"software\s+engineering", r"informatics"]
        ce_patterns = [r"computer\s+engineering", r"electrical\s+and\s+computer\s+engineering", r"\bece\b"]
        if any(re.search(p, t) for p in cs_patterns):
            degree_field = 'computer_science'
        elif any(re.search(p, t) for p in ce_patterns):
            degree_field = 'computer_engineering'
        
        return {'degree_level': degree_level, 'degree_field': degree_field}

class EnhancedFileReader:
    """Enhanced file reading with multiple fallback methods"""
    
    def __init__(self):
        self.text_processor = TextProcessor()
    
    def read_txt(self, path: Path) -> str:
        """Read text file with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding, errors='ignore')
            except UnicodeDecodeError:
                continue
        
        logger.warning(f"Could not decode text file: {path}")
        return ""
    
    def read_docx(self, path: Path) -> str:
        """Enhanced DOCX reading with better structure preservation"""
        try:
            doc = Document(str(path))
            content = []
            
            # Extract paragraphs with better formatting
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content.append(text)
            
            # Extract tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        content.append(row_text)
            
            return '\n'.join(content)
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {path}: {e}")
            return ""
    
    def read_pdf(self, path: Path) -> str:
        """Enhanced PDF reading with multiple extraction strategies"""
        text_parts = []
        
        try:
            # Primary method: pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        # Try text extraction
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                        else:
                            # Fallback: extract from tables if no text
                            tables = page.extract_tables()
                            for table in tables:
                                table_text = '\n'.join([' | '.join([str(cell) if cell else '' for cell in row]) for row in table])
                                text_parts.append(table_text)
                    
                    except Exception as e:
                        logger.warning(f"Error extracting page {i} from {path}: {e}")
                        continue
            
            extracted_text = '\n'.join(text_parts)
            
            # If extraction failed or produced minimal text, try alternative method
            if not extracted_text.strip() or len(extracted_text) < 100:
                logger.info(f"Trying alternative PDF extraction for {path}")
                extracted_text = self._try_textract(path)
            
            return self.text_processor.clean_text(extracted_text)
            
        except Exception as e:
            logger.error(f"Error reading PDF file {path}: {e}")
            return self._try_textract(path)
    
    def read_ods(self, path: Path) -> str:
        """Read ODS files using pyexcel-ods3 as specified in user rules"""
        try:
            data = pyexcel_ods3.get_data(str(path))
            content = []
            
            for sheet_name, sheet_data in data.items():
                content.append(f"Sheet: {sheet_name}")
                for row in sheet_data:
                    row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                    if row_text.strip():
                        content.append(row_text)
            
            return '\n'.join(content)
            
        except Exception as e:
            logger.error(f"Error reading ODS file {path}: {e}")
            return ""
    
    def _try_textract(self, path: Path) -> str:
        """Fallback text extraction using textract"""
        try:
            import textract
            return textract.process(str(path)).decode("utf-8", errors="ignore")
        except Exception as e:
            logger.warning(f"Textract failed for {path}: {e}")
            return ""
    
    def read_any(self, path: Path) -> Tuple[str, Dict[str, any]]:
        """Enhanced file reading with metadata extraction"""
        ext = path.suffix.lower()
        metadata = {
            'file_size': path.stat().st_size,
            'file_type': ext,
            'extraction_method': None
        }
        
        try:
            if ext == ".txt":
                text = self.read_txt(path)
                metadata['extraction_method'] = 'direct_text'
            elif ext == ".docx":
                text = self.read_docx(path)
                metadata['extraction_method'] = 'docx_python'
            elif ext == ".pdf":
                text = self.read_pdf(path)
                metadata['extraction_method'] = 'pdfplumber'
            elif ext == ".ods":
                text = self.read_ods(path)
                metadata['extraction_method'] = 'pyexcel_ods3'
            else:
                text = self._try_textract(path)
                metadata['extraction_method'] = 'textract_fallback'
            
            # Extract structured information
            structured_info = self.text_processor.extract_structured_info(text)
            metadata.update(structured_info)
            
            # Degree info
            degree_info = self.text_processor.detect_degree_info(text)
            metadata.update(degree_info)
            
            # Quality metrics
            metadata['char_count'] = len(text)
            metadata['word_count'] = len(text.split())
            metadata['line_count'] = len(text.split('\n'))
            
            return self.text_processor.clean_text(text), metadata
            
        except Exception as e:
            logger.error(f"Error reading file {path}: {e}")
            metadata['error'] = str(e)
            return "", metadata

class EnhancedKeywordAnalyzer:
    """Enhanced keyword analysis with fuzzy matching and context awareness"""
    
    def __init__(self):
        self.text_processor = TextProcessor()
    
    def normalize_columns(self, df: pd.DataFrame) -> pd.DataFrame:
        """Enhanced column normalization"""
        df = df.copy()
        # More comprehensive column name cleaning
        df.columns = (
            df.columns.astype(str)
                     .str.strip()
                     .str.lower()
                     .str.replace(r'\s+', '_', regex=True)
                     .str.replace(r'[^\w]', '_', regex=True)
                     .str.replace(r'_+', '_', regex=True)
                     .str.strip('_')
        )
        return df
    
    def detect_keyword_columns(self, cols: List[str]) -> Tuple[Optional[str], Optional[str]]:
        """Enhanced column detection with multiple patterns"""
        lower_cols = [c.lower() for c in cols]
        
        # Good keyword patterns
        good_patterns = ['good', 'positive', 'strong', 'relevant', 'desired']
        bad_patterns = ['bad', 'negative', 'weak', 'avoid', 'undesired']
        
        good_col = None
        bad_col = None
        
        for col in cols:
            col_lower = col.lower()
            if any(pattern in col_lower for pattern in good_patterns):
                if 'keyword' in col_lower or 'skill' in col_lower or 'term' in col_lower:
                    good_col = col
                    break
        
        for col in cols:
            col_lower = col.lower()
            if any(pattern in col_lower for pattern in bad_patterns):
                if 'keyword' in col_lower or 'skill' in col_lower or 'term' in col_lower:
                    bad_col = col
                    break
        
        return good_col, bad_col
    
    def split_and_normalize_terms(self, cell: str) -> List[str]:
        """Enhanced term splitting and normalization"""
        if pd.isna(cell) or not cell:
            return []
        
        # Split on various delimiters
        terms = re.split(r'[,;|\n\t]+', str(cell))
        
        normalized_terms = []
        for term in terms:
            # Clean and normalize each term
            term = term.strip().lower()
            term = re.sub(r'[^\w\s\-\+\#]', '', term)  # Remove special chars except useful ones
            term = re.sub(r'\s+', ' ', term)  # Normalize spaces
            
            if term and len(term) > 2:  # Filter out very short terms
                normalized_terms.append(term)
        
        return normalized_terms
    
    def create_enhanced_matcher(self, terms: List[str]) -> re.Pattern:
        """Create enhanced regex matcher with fuzzy matching capabilities"""
        if not terms:
            return re.compile(r'(?!)', re.IGNORECASE)  # Never matches
        
        # Sort by length (longest first) to match longer terms first
        sorted_terms = sorted(terms, key=len, reverse=True)
        
        # Escape special regex characters but preserve word boundaries
        escaped_terms = []
        for term in sorted_terms:
            if ' ' in term:  # Multi-word terms
                escaped_term = r'\b' + re.escape(term).replace(r'\ ', r'\s+') + r'\b'
            else:  # Single word terms
                escaped_term = r'\b' + re.escape(term) + r'\b'
            escaped_terms.append(escaped_term)
        
        pattern = '(' + '|'.join(escaped_terms) + ')'
        return re.compile(pattern, re.IGNORECASE | re.MULTILINE)
    
    def score_with_enhanced_keywords(self, df_texts: pd.DataFrame, kw_df: pd.DataFrame) -> pd.DataFrame:
        """Enhanced keyword scoring with context awareness"""
        kw_df = self.normalize_columns(kw_df)
        good_col, bad_col = self.detect_keyword_columns(list(kw_df.columns))
        
        if not good_col or not bad_col:
            available_cols = list(kw_df.columns)
            raise ValueError(
                f"Could not find keyword columns. Available columns: {available_cols}\n"
                f"Expected columns containing 'good' and 'bad' (or similar patterns)."
            )
        
        # Build comprehensive keyword sets
        all_good = set()
        all_bad = set()
        category_keywords = defaultdict(dict)
        
        for _, row in kw_df.iterrows():
            category = row.get('category', 'unknown')
            good_terms = self.split_and_normalize_terms(row[good_col])
            bad_terms = self.split_and_normalize_terms(row[bad_col])
            
            all_good.update(good_terms)
            all_bad.update(bad_terms)
            
            category_keywords[category] = {
                'good': good_terms,
                'bad': bad_terms
            }
        
        # Create enhanced matchers
        good_matcher = self.create_enhanced_matcher(sorted(all_good))
        bad_matcher = self.create_enhanced_matcher(sorted(all_bad))
        
        results = []
        for _, row in df_texts.iterrows():
            text = row['text'].lower()
            filename = row['filename']
            label = row['label']
            degree_level = row.get('degree_level', '') if isinstance(row, dict) else row.get('degree_level', '')
            degree_field = row.get('degree_field', '') if isinstance(row, dict) else row.get('degree_field', '')
            
            # Find matches
            good_matches = good_matcher.findall(text)
            bad_matches = bad_matcher.findall(text)
            
            # Count unique matches (case-insensitive)
            unique_good = list(set(match.lower() for match in good_matches))
            unique_bad = list(set(match.lower() for match in bad_matches))
            
            # Calculate scores by category
            category_scores = {}
            for category, keywords in category_keywords.items():
                cat_good = sum(1 for term in keywords['good'] if any(term in match for match in unique_good))
                cat_bad = sum(1 for term in keywords['bad'] if any(term in match for match in unique_bad))
                category_scores[f'{category}_good'] = cat_good
                category_scores[f'{category}_bad'] = cat_bad
            
            # Overall quality score
            total_good = len(unique_good)
            total_bad = len(unique_bad)
            quality_score = max(0, total_good - total_bad)  # Penalize bad keywords
            
            result = {
                'filename': filename,
                'label': label,
                'degree_level': degree_level,
                'degree_field': degree_field,
                'good_hits': total_good,
                'bad_hits': total_bad,
                'quality_score': quality_score,
                'good_terms_found': ', '.join(sorted(unique_good)),
                'bad_terms_found': ', '.join(sorted(unique_bad)),
                'word_count': len(text.split()),
                'char_count': len(text),
                **category_scores
            }
            
            results.append(result)
        
        return pd.DataFrame(results)

def load_resumes_enhanced(root: Path) -> pd.DataFrame:
    """Enhanced resume loading with better error handling and metadata"""
    reader = EnhancedFileReader()
    rows = []
    
    supported_extensions = {'.pdf', '.docx', '.txt', '.ods'}
    
    logger.info(f"Scanning for resumes in: {root}")
    
    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
            
        if file_path.suffix.lower() not in supported_extensions:
            continue
        
        try:
            text, metadata = reader.read_any(file_path)
            
            if not text or len(text.strip()) < 50:  # Minimum content threshold
                logger.warning(f"Insufficient content extracted from: {file_path}")
                continue
            
            # Determine label from folder structure
            label = file_path.parent.name
            if label == root.name:  # If file is directly in root
                label = 'general'
            
            row_data = {
                'filename': str(file_path.resolve()),
                'label': label,
                'text': text,
                **metadata
            }
            
            rows.append(row_data)
            logger.info(f"Processed: {file_path.name} ({metadata.get('word_count', 0)} words)")
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            continue
    
    df = pd.DataFrame(rows)
    
    if df.empty:
        logger.error("No valid resume files found!")
        return df
    
    logger.info(f"Successfully loaded {len(df)} resume files")
    return df

def build_enhanced_bow(df_texts: pd.DataFrame, 
                      stopwords: Union[str, List[str]] = "english",
                      max_features: int = 8000,
                      ngrams: Tuple[int, int] = (1, 2)) -> Tuple[pd.DataFrame, CountVectorizer]:
    """Enhanced bag-of-words with better preprocessing"""
    
    # Custom stopwords for resume analysis
    resume_stopwords = {
        'university', 'college', 'school', 'student', 'education',
        'resume', 'cv', 'curriculum', 'vitae', 'experience', 'work',
        'email', 'phone', 'address', 'linkedin', 'github'
    }
    
    if isinstance(stopwords, str) and stopwords == "english":
        stop_words = set(nltk.corpus.stopwords.words('english'))
        stop_words.update(resume_stopwords)
        stopwords = list(stop_words)
    elif isinstance(stopwords, list):
        stopwords.extend(list(resume_stopwords))
    
    # Enhanced vectorizer with better parameters
    vectorizer = CountVectorizer(
        stop_words=stopwords,
        max_features=max_features,
        ngram_range=ngrams,
        min_df=2,  # Ignore terms that appear in less than 2 documents
        max_df=0.8,  # Ignore terms that appear in more than 80% of documents
        lowercase=True,
        token_pattern=r'\b[a-zA-Z][a-zA-Z0-9\-\+\#]*[a-zA-Z0-9]\b|\b[a-zA-Z]\b'  # Better token pattern
    )
    
    try:
        X = vectorizer.fit_transform(df_texts['text'])
        feature_names = vectorizer.get_feature_names_out()
        
        bow = pd.DataFrame(
            X.toarray(), 
            index=df_texts['filename'], 
            columns=feature_names
        )
        
        logger.info(f"Created BOW matrix: {bow.shape[0]} documents × {bow.shape[1]} features")
        return bow, vectorizer
        
    except Exception as e:
        logger.error(f"Error creating BOW matrix: {e}")
        raise

def write_default_keywords_csv(path: Path, overwrite: bool = False) -> None:
    """Create a curated default keywords CSV covering CS and CE domains."""
    if path.exists() and not overwrite:
        return
    rows = [
        {"Category": "Programming", "Good_Keywords": "python, java, c, c++, c#, javascript, typescript, go, rust, kotlin, swift, matlab, r", "Bad_Keywords": "basic computer skills, microsoft office only, beginner"},
        {"Category": "Data Science", "Good_Keywords": "machine learning, deep learning, nlp, ai, pandas, numpy, scikit-learn, tensorflow, pytorch, statistics, data mining, computer vision", "Bad_Keywords": "some experience with data, basic statistics"},
        {"Category": "Cloud/DevOps", "Good_Keywords": "aws, azure, gcp, docker, kubernetes, ci/cd, terraform, ansible, jenkins, linux, bash", "Bad_Keywords": "familiar with cloud, tinkered with aws"},
        {"Category": "Tools", "Good_Keywords": "git, github, gitlab, jira, confluence, tableau, power bi, spark, hadoop, postgres, mysql, sqlite, mongodb", "Bad_Keywords": "microsoft paint, office only"},
        {"Category": "Soft Skills", "Good_Keywords": "leadership, teamwork, communication, collaboration, problem solving", "Bad_Keywords": "hardworking, go-getter, fast learner"},
        {"Category": "Certifications", "Good_Keywords": "aws certified, azure certified, gcp certified, pmp, scrum master", "Bad_Keywords": "certificate of participation, attended workshop"},
        {"Category": "Experience Level", "Good_Keywords": "senior, lead, manager, 5+ years, expert", "Bad_Keywords": "entry-level, beginner, novice"},
        {"Category": "Impact Words", "Good_Keywords": "delivered, achieved, designed, implemented, optimized, led, built, created, improved, reduced", "Bad_Keywords": "helped with, tried, attempted, exposed to"},
        {"Category": "Computer Science Core", "Good_Keywords": "data structures, algorithms, operating systems, computer architecture, discrete mathematics, databases, networking, software engineering, object-oriented programming, compilers", "Bad_Keywords": "basic computer skills"},
        {"Category": "Computer Engineering Core", "Good_Keywords": "embedded systems, digital logic, verilog, vhdl, fpga, microcontroller, microcontrollers, microprocessor, systemverilog, pcb design, circuits, hardware design, signal processing, control systems", "Bad_Keywords": "breadboard only, simple circuits only"}
    ]
    pd.DataFrame(rows).to_csv(path, index=False)
    logger.info(f"Wrote default keywords CSV → {path}")

def main():
    """Enhanced main function with comprehensive analysis pipeline"""
    
    parser = argparse.ArgumentParser(
        description="Enhanced Resume Analysis - Systematic data extraction and analysis",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument("--root", required=True, 
                       help="Root folder containing resume files (PDF/DOCX/TXT/ODS)")
    parser.add_argument("--outdir", default="out", 
                       help="Output directory for analysis results")
    parser.add_argument("--keywords", default="resume_keywords.csv",
                       help="CSV file with keyword categories and terms")
    parser.add_argument("--overwrite-keywords", action="store_true",
                       help="Overwrite keywords file with defaults")
    parser.add_argument("--min-content-length", type=int, default=100,
                       help="Minimum content length for valid resumes")
    parser.add_argument("--max-features", type=int, default=8000,
                       help="Maximum features for BOW analysis")
    parser.add_argument("--debug", action="store_true",
                       help="Enable debug logging")
    
    args = parser.parse_args()
    
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Setup paths
    root_path = Path(args.root).resolve()
    output_dir = Path(args.outdir).resolve()
    keywords_path = Path(args.keywords).resolve()
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Ensure keywords file exists
    if not keywords_path.exists() or args.overwrite_keywords:
        logger.info("Creating default keywords file...")
        write_default_keywords_csv(keywords_path, overwrite=True)
    
    logger.info("=" * 50)
    logger.info("ENHANCED RESUME ANALYSIS PIPELINE")
    logger.info("=" * 50)
    
    try:
        # 1. Load and process resumes
        logger.info("Step 1: Loading resumes...")
        df_texts = load_resumes_enhanced(root_path)
        
        if df_texts.empty:
            logger.error("No resumes loaded. Exiting.")
            sys.exit(1)
        
        # Save processed texts with metadata
        texts_output = output_dir / "resume_texts_enhanced.csv"
        df_texts.to_csv(texts_output, index=False)
        logger.info(f"Saved enhanced text data: {texts_output}")
        
        # 2. Build enhanced BOW matrix
        logger.info("Step 2: Building enhanced BOW matrix...")
        bow, vectorizer = build_enhanced_bow(df_texts, max_features=args.max_features)
        
        bow_output = output_dir / "resume_word_counts_enhanced.csv"
        bow.to_csv(bow_output)
        logger.info(f"Saved enhanced BOW matrix: {bow_output}")
        
        # 3. Enhanced keyword analysis
        if keywords_path.exists():
            logger.info("Step 3: Enhanced keyword analysis...")
            analyzer = EnhancedKeywordAnalyzer()
            kw_df = pd.read_csv(keywords_path)
            
            scored = analyzer.score_with_enhanced_keywords(df_texts, kw_df)
            scored_output = output_dir / "resume_keyword_scores_enhanced.csv"
            scored.to_csv(scored_output, index=False)
            logger.info(f"Saved enhanced keyword scores: {scored_output}")
        
        # 4. Enhanced similarity analysis
        logger.info("Step 4: Enhanced similarity analysis...")
        tfidf = TfidfVectorizer(
            stop_words='english',
            max_features=5000,
            ngram_range=(1, 2),
            min_df=2,
            max_df=0.8
        )
        
        X = tfidf.fit_transform(df_texts['text'])
        similarity_matrix = cosine_similarity(X)
        
        sim_df = pd.DataFrame(
            similarity_matrix,
            index=df_texts['filename'],
            columns=df_texts['filename']
        )
        
        sim_output = output_dir / "resume_similarity_enhanced.csv"
        sim_df.to_csv(sim_output)
        logger.info(f"Saved enhanced similarity matrix: {sim_output}")
        
        logger.info("=" * 50)
        logger.info("ANALYSIS COMPLETE!")
        logger.info(f"Results saved in: {output_dir}")
        logger.info("=" * 50)
        
    except Exception as e:
        logger.error(f"Analysis failed: {e}")
        raise

if __name__ == "__main__":
    main()
